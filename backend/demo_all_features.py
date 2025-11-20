import os
import json
import asyncio
import logging
from datetime import datetime

# Setup logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Import our modules
from app.vector_store.factory import VectorStoreFactory
from app.document_management.generators.template_renderer import TemplateRenderer
from app.document_management.generators.pdf_generator import PDFGenerator
from app.document_management.generators.word_generator import WordGenerator
from app.document_management.storage.local_storage import LocalStorage
from app.document_management.security.encryption import DocumentEncryption
from app.document_management.security.access_control import AccessControlManager

async def main():
    print("="*50)
    print("       FEATURE DEMONSTRATION SCRIPT       ")
    print("="*50)

    # ---------------------------------------------------------
    # 1. VECTOR STORE DEMONSTRATION
    # ---------------------------------------------------------
    logger.info("--- 1. Vector Store Demonstration (ChromaDB) ---")

    # Configuration for ChromaDB (Local)
    chroma_config = {
        "embedding_dimension": 384,
        "collection_name": "demo_collection",
        "persist_directory": "./demo_vector_store",
        "embedding_model": "all-MiniLM-L6-v2"
    }

    # Initialize Store
    try:
        vector_store = VectorStoreFactory.create_store("chroma", chroma_config)

        # Add Data
        texts = [
            "The quarterly financial report shows a 20% increase in revenue.",
            "Employee handbook: working hours are 9 AM to 5 PM.",
            "Project X is a confidential initiative to build a new AI model."
        ]
        ids = ["doc1", "doc2", "doc3"]
        metadata = [
            {"type": "financial", "year": 2023},
            {"type": "hr", "public": True},
            {"type": "project", "confidential": True}
        ]

        logger.info(f"Adding {len(texts)} documents to vector store...")
        vector_store.add_texts(texts, ids, metadata)

        # Search
        query = "revenue growth"
        logger.info(f"Searching for: '{query}'")
        results = vector_store.search(query, top_k=1)

        for res in results:
            print(f"   Found: {res.raw_content} (Score: {res.similarity_score:.4f})")

        # Stats
        stats = vector_store.get_stats()
        logger.info(f"Store Stats: {stats}")

    except Exception as e:
        logger.error(f"Vector store demo failed: {e}")

    print("\n")

    # ---------------------------------------------------------
    # 2. DOCUMENT GENERATION DEMONSTRATION
    # ---------------------------------------------------------
    logger.info("--- 2. Document Generation Demonstration ---")

    # Data for templates
    context = {
        "company_name": "Acme Corp",
        "report_date": datetime.now().strftime("%Y-%m-%d"),
        "revenue": "$1,200,000",
        "growth": "20%",
        "items": [
            {"name": "Product A", "sales": 5000},
            {"name": "Product B", "sales": 7000}
        ]
    }

    # A. PDF Generation
    logger.info("Generating PDF Report...")
    html_template = """
    <html>
    <head>
        <style>
            body { font-family: Helvetica; }
            h1 { color: #333; }
            table { width: 100%; border-collapse: collapse; }
            th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
            th { background-color: #f2f2f2; }
        </style>
    </head>
    <body>
        <h1>Quarterly Report - {{ company_name }}</h1>
        <p>Date: {{ report_date }}</p>
        <p>Total Revenue: <strong>{{ revenue }}</strong> (Growth: {{ growth }})</p>

        <h2>Sales Breakdown</h2>
        <table>
            <tr><th>Product</th><th>Sales</th></tr>
            {% for item in items %}
            <tr><td>{{ item.name }}</td><td>{{ item.sales }}</td></tr>
            {% endfor %}
        </table>
    </body>
    </html>
    """

    rendered_html = TemplateRenderer.render(html_template, context)
    pdf_bytes = PDFGenerator.generate_from_html(rendered_html)
    logger.info(f"PDF Generated ({len(pdf_bytes)} bytes)")

    # Watermark
    logger.info("Adding Watermark...")
    watermarked_pdf = PDFGenerator.add_watermark(pdf_bytes, "CONFIDENTIAL")
    logger.info(f"Watermark added. Size: {len(watermarked_pdf)} bytes")

    # B. Word Generation
    logger.info("Generating Word Contract...")
    # Create a quick dummy docx template in memory for demo purposes
    from docx import Document
    import io

    # Create template
    tpl_doc = Document()
    tpl_doc.add_heading('Contract Agreement', 0)
    tpl_doc.add_paragraph('This agreement is between {{ company_name }} and the Client.')
    tpl_doc.add_paragraph('Dated: {{ report_date }}')

    tpl_buffer = io.BytesIO()
    tpl_doc.save(tpl_buffer)
    tpl_bytes = tpl_buffer.getvalue()

    # Generate
    word_bytes = WordGenerator.generate(tpl_bytes, context)
    logger.info(f"Word Document Generated ({len(word_bytes)} bytes)")

    print("\n")

    # ---------------------------------------------------------
    # 3. SECURITY DEMONSTRATION
    # ---------------------------------------------------------
    logger.info("--- 3. Security Demonstration ---")

    # Encrypt PDF
    logger.info("Encrypting PDF...")
    encrypted_pdf = DocumentEncryption.encrypt_pdf(
        watermarked_pdf,
        user_password="password123",
        owner_password="admin"
    )
    logger.info(f"PDF Encrypted. Size: {len(encrypted_pdf)} bytes")

    # Protect Word
    logger.info("Protecting Word Document (Restrict Editing)...")
    protected_word = DocumentEncryption.encrypt_docx(word_bytes, "password")
    logger.info(f"Word Document Protected. Size: {len(protected_word)} bytes")

    # Access Control
    logger.info("Checking Access Control...")
    rbac = AccessControlManager()

    user_role = "user"
    admin_role = "admin"

    can_download = rbac.check_permission(user_role, "download")
    can_delete = rbac.check_permission(user_role, "delete")
    logger.info(f"Role '{user_role}' can download? {can_download}")
    logger.info(f"Role '{user_role}' can delete? {can_delete}")

    can_admin_delete = rbac.check_permission(admin_role, "delete")
    logger.info(f"Role '{admin_role}' can delete? {can_admin_delete}")

    print("\n")

    # ---------------------------------------------------------
    # 4. STORAGE DEMONSTRATION
    # ---------------------------------------------------------
    logger.info("--- 4. Storage Demonstration (Local) ---")

    storage = LocalStorage("./demo_storage")

    # Save files
    logger.info("Saving files to ./demo_storage/ ...")
    await storage.upload("reports/report_2023_Q3.pdf", encrypted_pdf)
    await storage.upload("contracts/contract_acme.docx", protected_word)

    exists_pdf = await storage.exists("reports/report_2023_Q3.pdf")
    exists_word = await storage.exists("contracts/contract_acme.docx")

    logger.info(f"PDF stored successfully? {exists_pdf}")
    logger.info(f"Word stored successfully? {exists_word}")

    # Set Readonly
    await storage.set_readonly("reports/report_2023_Q3.pdf", True)
    logger.info("Marked PDF as Read-Only in filesystem.")

    print("="*50)
    print("       DEMONSTRATION COMPLETED       ")
    print("="*50)

if __name__ == "__main__":
    # Ensure we are running from backend directory context if needed,
    # but imports assume we are in root of repo or pythonpath set.
    # Let's try running it.
    asyncio.run(main())
