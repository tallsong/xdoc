from docx import Document
from docx.shared import Inches
import io
from typing import Dict, Any
import re

class WordGenerator:
    @staticmethod
    def generate(template_content: bytes, data: Dict[str, Any]) -> bytes:
        """
        Generate Word document by filling placeholders in a template.
        Preserves formatting by iterating over runs.
        """
        doc = Document(io.BytesIO(template_content))

        def replace_text_in_paragraph(paragraph, data):
            # We need to handle cases where {{ key }} is split across runs.
            # A simple approach is to iterate over runs and check if any run contains the full placeholder.
            # If so, replace it.
            # If the placeholder is split, this simple logic won't work without complex run merging.
            # For this implementation, we will try to replace within runs first.
            # If we can't find it in runs but it is in paragraph.text, we fall back to paragraph.text replacement
            # (losing formatting for that specific paragraph), OR we assume placeholders are typed in one go.

            # Improved strategy:
            # Check if paragraph contains any placeholder
            has_placeholder = False
            for key in data.keys():
                if f"{{{{{key}}}}}" in paragraph.text:
                    has_placeholder = True
                    break

            if not has_placeholder:
                return

            # Try to replace in runs
            replaced_in_runs = False
            for run in paragraph.runs:
                for key, value in data.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(value))
                        replaced_in_runs = True

            # If we still have placeholders in paragraph text but couldn't replace in individual runs
            # (meaning it was split across runs), we are forced to clobber runs to replace text
            # to ensure correctness of data over formatting in that specific edge case.
            # But we can try to preserve the style of the first run.
            if not replaced_in_runs:
                 # check again if still has placeholder (maybe some were replaced)
                still_has_placeholder = False
                for key in data.keys():
                    if f"{{{{{key}}}}}" in paragraph.text:
                         still_has_placeholder = True
                         break

                if still_has_placeholder:
                    # Fallback: replace paragraph text.
                    # Attempt to keep style of first run if exists
                    style = None
                    if paragraph.runs:
                        style = paragraph.runs[0].style

                    text = paragraph.text
                    for key, value in data.items():
                         text = text.replace(f"{{{{{key}}}}}", str(value))

                    paragraph.text = text # This clears runs
                    if style and paragraph.runs:
                        paragraph.runs[0].style = style


        # Iterate over paragraphs
        for paragraph in doc.paragraphs:
            replace_text_in_paragraph(paragraph, data)

        # Iterate over tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                         replace_text_in_paragraph(paragraph, data)

        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()
