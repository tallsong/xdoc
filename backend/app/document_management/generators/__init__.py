"""Document generators for PDF and Word formats."""

import logging
from abc import ABC, abstractmethod
from typing import Any, Dict, Optional, List
from datetime import datetime
from io import BytesIO
import json

logger = logging.getLogger(__name__)


class DocumentGenerator(ABC):
    """Abstract document generator."""

    @abstractmethod
    def generate(
        self,
        template_content: bytes,
        data: Dict[str, Any],
        **options,
    ) -> bytes:
        """Generate document from template and data.

        Args:
            template_content: Template file bytes
            data: Data to fill into template
            **options: Generator-specific options

        Returns:
            Generated document bytes
        """
        pass


class PDFGenerator(DocumentGenerator):
    """PDF document generator."""

    @staticmethod
    def generate_from_html(
        html_content: str,
        css: Optional[str] = None,
    ) -> bytes:
        """Generate PDF from HTML.

        Args:
            html_content: HTML content
            css: CSS styles

        Returns:
            PDF bytes
        """
        try:
            from weasyprint import HTML, CSS
            from io import BytesIO

            logger.info("Generating PDF from HTML using WeasyPrint")

            doc = HTML(string=html_content, base_url=".")

            if css:
                style = CSS(string=css)
                pdf_bytes = doc.write_pdf(stylesheets=[style])
            else:
                pdf_bytes = doc.write_pdf()

            logger.info(f"PDF generated successfully (size: {len(pdf_bytes)})")
            return pdf_bytes

        except ImportError:
            logger.error("weasyprint not installed")
            raise ImportError("weasyprint required for PDF generation")

    @staticmethod
    def add_watermark(
        pdf_content: bytes,
        watermark_text: str = "Internal Use Only",
        opacity: float = 0.3,
    ) -> bytes:
        """Add watermark to PDF.

        Args:
            pdf_content: PDF bytes
            watermark_text: Watermark text
            opacity: Watermark opacity (0-1)

        Returns:
            PDF with watermark
        """
        try:
            import PyPDF2
            from io import BytesIO
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.colors import HexColor

            logger.info(f"Adding watermark: {watermark_text}")

            # Create watermark
            watermark_buffer = BytesIO()
            c = canvas.Canvas(watermark_buffer, pagesize=letter)
            c.setFillAlpha(opacity)
            c.setFont("Helvetica", 60)
            c.rotate(45)
            c.drawString(200, 100, watermark_text)
            c.save()

            watermark_buffer.seek(0)
            watermark_pdf = PyPDF2.PdfReader(watermark_buffer)
            watermark_page = watermark_pdf.pages[0]

            # Apply watermark to all pages
            reader = PyPDF2.PdfReader(BytesIO(pdf_content))
            writer = PyPDF2.PdfWriter()

            for page in reader.pages:
                page.merge_page(watermark_page)
                writer.add_page(page)

            output = BytesIO()
            writer.write(output)
            logger.info("Watermark added successfully")
            return output.getvalue()

        except Exception as e:
            logger.error(f"Watermark addition failed: {e}")
            # Return original if watermarking fails
            return pdf_content


class WordGenerator(DocumentGenerator):
    """Word document generator."""

    def generate(
        self,
        template_content: bytes,
        data: Dict[str, Any],
        **options,
    ) -> bytes:
        """Generate Word document from template and data.

        Args:
            template_content: Template DOCX bytes
            data: Data to fill into template
            **options: Options like tables_data, images

        Returns:
            Generated DOCX bytes
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from io import BytesIO
            import re

            logger.info("Generating Word document from template")

            doc = Document(BytesIO(template_content))

            # Process paragraphs
            for paragraph in doc.paragraphs:
                self._process_paragraph(paragraph, data)

            # Process tables
            for table in doc.tables:
                self._process_table(table, data)

            # Add dynamic tables if provided
            if "tables_data" in options:
                self._add_dynamic_tables(doc, options["tables_data"])

            # Add images if provided
            if "images" in options:
                self._add_images(doc, options["images"])

            output = BytesIO()
            doc.save(output)
            logger.info("Word document generated successfully")
            return output.getvalue()

        except ImportError:
            logger.error("python-docx not installed")
            raise ImportError("python-docx required for Word generation")

    @staticmethod
    def _process_paragraph(paragraph, data: Dict[str, Any]) -> None:
        """Process paragraph text with data substitution."""
        text = paragraph.text
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"  # {{key}}
            if placeholder in text:
                text = text.replace(placeholder, str(value))

        if text != paragraph.text:
            paragraph.text = text

    @staticmethod
    def _process_table(table, data: Dict[str, Any]) -> None:
        """Process table cells with data substitution."""
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    WordGenerator._process_paragraph(paragraph, data)

    @staticmethod
    def _add_dynamic_tables(doc, tables_data: List[Dict[str, Any]]) -> None:
        """Add dynamic tables to document.

        Args:
            doc: Document object
            tables_data: List of table data dicts with keys: headers, rows
        """
        for table_config in tables_data:
            # Add table
            rows = len(table_config.get("rows", [])) + 1  # +1 for header
            cols = len(table_config.get("headers", []))

            table = doc.add_table(rows=rows, cols=cols)
            table.style = "Light Grid Accent 1"

            # Add headers
            header_cells = table.rows[0].cells
            for i, header in enumerate(table_config.get("headers", [])):
                header_cells[i].text = str(header)

            # Add rows
            for row_idx, row_data in enumerate(table_config.get("rows", []), start=1):
                row_cells = table.rows[row_idx].cells
                for col_idx, value in enumerate(row_data):
                    row_cells[col_idx].text = str(value)

    @staticmethod
    def _add_images(doc, images: Dict[str, bytes]) -> None:
        """Add images to document.

        Args:
            doc: Document object
            images: Dict of {field_name: image_bytes}
        """
        from io import BytesIO

        for field_name, image_bytes in images.items():
            try:
                doc.add_paragraph(f"[{field_name}]")
                image_stream = BytesIO(image_bytes)
                doc.add_picture(image_stream, width=Inches(4))
            except Exception as e:
                logger.error(f"Failed to add image {field_name}: {e}")


class HTMLGenerator(DocumentGenerator):
    """HTML document generator."""

    def generate(
        self,
        template_content: bytes,
        data: Dict[str, Any],
        **options,
    ) -> bytes:
        """Generate HTML document from template and data.

        Args:
            template_content: Template HTML bytes
            data: Data to fill into template
            **options: Options like css

        Returns:
            Generated HTML bytes
        """
        try:
            from jinja2 import Template

            logger.info("Generating HTML document")

            template_str = template_content.decode("utf-8")
            template = Template(template_str)
            html = template.render(**data)

            logger.info("HTML document generated successfully")
            return html.encode("utf-8")

        except Exception as e:
            logger.error(f"HTML generation failed: {e}")
            raise


class TemplateRenderer:
    """Render templates with data using Jinja2."""

    @staticmethod
    def render(template_str: str, data: Dict[str, Any]) -> str:
        """Render template string with data.

        Args:
            template_str: Template string with placeholders
            data: Data to render

        Returns:
            Rendered string
        """
        try:
            from jinja2 import Template, TemplateError

            logger.debug("Rendering template with data")

            template = Template(template_str)
            result = template.render(**data)

            return result

        except TemplateError as e:
            logger.error(f"Template rendering failed: {e}")
            raise

    @staticmethod
    def render_json(json_str: str, data: Dict[str, Any]) -> Dict[str, Any]:
        """Render JSON template with data.

        Args:
            json_str: JSON template string
            data: Data to render

        Returns:
            Rendered dictionary
        """
        # First render the JSON string as a template
        rendered_str = TemplateRenderer.render(json_str, data)
        # Then parse as JSON
        return json.loads(rendered_str)
